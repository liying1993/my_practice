"""bookify.py: Combines multiple Word docs in a folder.

"""

import os
import sys
from glob import glob

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Inches, Pt
except ImportError:
    raise ImportError("You need to 'pip install python-docx'")

TEXT_FONT = "Trebuchet MS"


def add_matter(master_document, filename, chapter, after=False):
    """Builds """
    if not os.path.exists(filename):
        return master_document

    if after:
        master_document.add_page_break()

    # Build the heading
    heading = master_document.add_heading('', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runt = heading.add_run(chapter)
    runt.font.color.theme_color = WD_COLOR_INDEX.WHITE

    # Add the material
    document = Document(docx=filename)
    for index, paragraph in enumerate(document.paragraphs):
        new_paragraph = master_document.add_paragraph()
        new_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
        new_paragraph.style = paragraph.style
        # Loop through the runs of a paragraph
        # A run is a style element within a paragraph (i.e. bold)
        for j, run in enumerate(paragraph.runs):
            # Copy over the old style
            text = run.text
            # Add run to new paragraph
            new_run = new_paragraph.add_run(text=text)
            # Update styles for run
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.font.size = run.font.size
            new_run.font.color.theme_color = WD_COLOR_INDEX.BLACK
    master_document.add_page_break()
    print(f'Adding {chapter}')
    return master_document


def add_chapter(master_document, filename, chapter):
    """Build chapters, i.e. where the story happens."""
    # Build the chapter
    document = Document(docx=filename)

    # Build the heading
    heading = master_document.add_heading('', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading.add_run(chapter).font.color.theme_color = WD_COLOR_INDEX.BLACK
    heading.paragraph_format.space_after = Pt(12)

    for index, paragraph in enumerate(document.paragraphs):
        new_paragraph = master_document.add_paragraph()
        # Loop through the runs of a paragraph
        # A run is a style element within a paragraph (i.e. bold)
        for j, run in enumerate(paragraph.runs):

            text = run.text
            # If at start of paragraph and no tab, add one
            if j == 0 and not text.startswith('\t'):
                text = f"\t{text}"
            # Add run to new paragraph
            new_run = new_paragraph.add_run(text=text)
            # Update styles for run
            new_run.font.name = TEXT_FONT
            new_run.bold = run.bold
            new_run.italic = run.italic

        # Last minute format checking
        text = new_paragraph.text

    master_document.add_page_break()
    # Destroy the document object
    del document
    return master_document


def main(book):
    master_document = Document()

    master_document = add_matter(
      master_document,
      filename=f'{book}/_title.docx',
      chapter='Title Page'
    )
    master_document = add_matter(
        master_document,
        filename=f'{book}/_copyright.docx',
        chapter='Copyright Page'
    )
    master_document = add_matter(
        master_document,
        filename=f'{book}/_dedication.docx',
        chapter='Dedication'
    )

    for filename in glob(f"{book}/*"):
        if filename.startswith(f"{book}/_"):
            print(f'skipping {filename}')
            continue

        # Get the chapter name
        book, short = filename.split('/')
        chapter = short.replace('.docx', '')
        if chapter.startswith('0'):
            chapter = chapter[1:]
        print(f'Adding {chapter}')
        master_document = add_chapter(master_document, filename, chapter)

    master_document = add_matter(
        master_document,
        filename=f'{book}/_aboutauthor.docx',
        chapter='About the Author',
        after=True
    )
    master_document = add_matter(
        master_document,
        filename=f'{book}/_afterward.docx',
        chapter='Afterward',
        after=True
    )
    master_document.save(f'{book}.docx')
    print('DONE!!!')

if __name__ == '__main__':
    try:
        book = sys.argv[1]
    except IndexError:
        msg = 'You need to specify a book. A book is a directory of word files.'
        raise Exception(msg)

    main(book)